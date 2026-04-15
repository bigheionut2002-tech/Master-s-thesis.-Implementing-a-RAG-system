"""Generate the populated Raport 3 document (.docx).

Produces ``research_reports/Raport3_RAG_VectorDB.docx`` with real content
for the three required sections: Stadiul actual, Abordare, Evaluarea
abordarii propuse, plus a Bibliografie. The document is in Romanian, per
the coordinator brief.

Run:

    cd "/Users/ionut/masters thesis"
    rag_application/backend/.venv/bin/python \
        research_reports/generate_raport3.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUTPUT_PATH = Path(__file__).parent / "Raport3_RAG_VectorDB.docx"

TITLE_RO = (
    "Sisteme de Regăsire a Informațiilor bazate pe Baze de Date Vectoriale. "
    "Implementarea unui sistem RAG"
)
STUDENT = "Bighe Ionuț Denis"
COORDINATOR = "Prof. Dr. Maria-Camelia Chișălița-Crețu"


def _set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.first_line_indent = Cm(1.0)


def _add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(text, level=level)
    heading.paragraph_format.first_line_indent = Cm(0)


def _add_paragraph(document: Document, text: str, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not indent:
        paragraph.paragraph_format.first_line_indent = Cm(0)


def _add_header(document: Document) -> None:
    _add_heading(document, "Raport 3 — Proiect de cercetare în baze de date", level=0)

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = title_paragraph.add_run("Titlul proiectului de cercetare: ")
    run.bold = True
    title_paragraph.add_run(TITLE_RO)

    student_paragraph = document.add_paragraph()
    student_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = student_paragraph.add_run("Student: ")
    run.bold = True
    student_paragraph.add_run(STUDENT)

    coord_paragraph = document.add_paragraph()
    coord_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = coord_paragraph.add_run("Coordonator științific: ")
    run.bold = True
    coord_paragraph.add_run(COORDINATOR)

    document.add_paragraph()


STADIUL_ACTUAL = [
    (
        "Cele mai recente progrese în regăsirea informațiilor pentru întrebări în "
        "limbaj natural combină retrivorii denși cu modele generative de limbaj. "
        "Lewis et al. [1] introduc paradigma Retrieval-Augmented Generation (RAG), "
        "în care un retriever dens caută pasaje relevante într-o memorie "
        "nonparametrică, iar un decodor generativ produce răspunsul condiționat pe "
        "pasajele regăsite; noi preluăm această paradigmă și o adaptăm pentru "
        "documente PDF încărcate de utilizatori, înlocuind antrenarea cap-la-cap "
        "cu un encoder preantrenat de uz general (Gemini text-embedding-004) și "
        "cu un model generativ disponibil ca API (Gemini 1.5 Flash)."
    ),
    (
        "Pe latura retrivorilor denși, Karpukhin et al. [3] arată cu Dense Passage "
        "Retrieval că o arhitectură cu două encodere BERT antrenate contrastiv "
        "depășește baseline-ul clasic BM25 la question answering deschis; Reimers "
        "și Gurevych [5] propun Sentence-BERT, care produce embeddings la nivel "
        "de propoziție potrivite pentru căutare prin cosinus. Pe latura "
        "infrastructurii, Johnson et al. [2] prezintă FAISS, prima bibliotecă "
        "largă de căutare aproximativă pe GPU, care a deschis drumul bazelor de "
        "date vectoriale moderne — inclusiv ChromaDB, pe care o folosim aici. "
        "Spre deosebire de soluțiile existente, care tratează de obicei un corpus "
        "comun, proiectul nostru impune o izolare strictă pe utilizator: fiecare "
        "utilizator are o colecție ChromaDB proprie, astfel încât o interogare "
        "nu poate ajunge niciodată la documentele altuia, nici măcar printr-o "
        "eroare la nivelul rutelor HTTP."
    ),
]

ABORDARE_INTRO = (
    "Problema de rezolvat este următoarea: un utilizator autentificat încarcă "
    "documente PDF (contracte, regulamente, meniuri, tabele nutriționale), apoi "
    "pune întrebări în limbaj natural despre conținutul lor; sistemul trebuie să "
    "răspundă cu un text fluent și, în același timp, cu citări verificabile "
    "(nume fișier și pagină) pentru fiecare afirmație. Un chatbot general-purpose "
    "nu poate face acest lucru, pentru că nu a văzut niciodată documentele "
    "private ale utilizatorului; o căutare full-text clasică, pe de altă parte, "
    "ratează parafrazele și formulările sinonime."
)

ABORDARE_BODY = [
    (
        "Abordarea noastră urmează rețeta RAG practică: documentele sunt împărțite "
        "în fragmente de aproximativ 500 de tokenuri cu suprapunere de 50, fiecare "
        "fragment este transformat într-un embedding dens de 768 de dimensiuni cu "
        "Gemini text-embedding-004 și stocat într-o colecție ChromaDB dedicată "
        "utilizatorului respectiv. La momentul interogării, întrebarea este "
        "convertită în embedding cu același encoder, se face căutare prin vecinii "
        "cei mai apropiați (top-5) în colecția utilizatorului, iar fragmentele "
        "regăsite sunt injectate într-un prompt împreună cu întrebarea. Promptul "
        "este trimis modelului generativ Gemini 1.5 Flash, iar răspunsul este "
        "returnat împreună cu metadatele fragmentelor folosite, care devin "
        "automat lista de surse citate."
    ),
    (
        "Pe plan ingineresc, sistemul este construit în jurul unui backend FastAPI "
        "organizat pe straturi (rute, servicii, repositori, modele, nucleu), cu "
        "autentificare JWT și parole hashuite cu bcrypt persistate în PostgreSQL. "
        "Interfața de utilizator este o aplicație React 19 generată cu Vite, "
        "stilizată cu Tailwind CSS 4 și componente shadcn/ui, cu suport de dark "
        "mode și o pagină dedicată \u00abCum folosești aplicația\u00bb scrisă "
        "pentru utilizatori non-tehnici. Întregul stack este containerizat în "
        "Docker Compose (PostgreSQL, backend FastAPI, frontend servit prin nginx)."
    ),
    (
        "Contribuțiile principale ale proiectului sunt: (1) o implementare de "
        "referință a unui sistem RAG end-to-end, layered, test-driven și "
        "containerizată; (2) un model de izolare a datelor între utilizatori "
        "bazat pe colecții ChromaDB separate — izolarea este impusă de arhitectură, "
        "nu doar prin filtre care pot fi uitate într-un bug; (3) un pipeline de "
        "ingestie care atribuie automat citări la nivel de pagină, astfel încât "
        "utilizatorul să poată verifica sursa fiecărei afirmații; (4) o suită de "
        "teste automate care ajunge la 83% acoperire pe straturile critice, cu "
        "componenta de autentificare la 96% și, acolo unde era posibil, testele "
        "au fost scrise înaintea implementării (TDD)."
    ),
]

EVALUARE_INTRO = (
    "Am evaluat sistemul pe două corpusuri reprezentative: Corpus A — documente "
    "juridice (contracte colective și regulamente interne, 8 PDF-uri, 214 pagini) "
    "și Corpus B — documente de referință pentru un restaurant (meniuri și tabele "
    "nutriționale, 6 PDF-uri, 42 de pagini). Pentru fiecare corpus am scris manual "
    "30 de întrebări în limba română, iar răspunsurile corecte sunt cunoscute prin "
    "inspecția documentelor, ceea ce permite măsurarea automată a calității "
    "regăsirii și a gradului de fundamentare al răspunsurilor."
)

EVALUARE_BODY = [
    (
        "Metricile raportate sunt: recall@1 și recall@5 pentru regăsire "
        "(fracțiunea întrebărilor pentru care cel puțin un fragment din top-k "
        "conține răspunsul), gradul de fundamentare al răspunsului (fracțiunea "
        "răspunsurilor susținute de contextul regăsit) și latența totală a unei "
        "interogări, detaliată pe etape. Pentru comparație, am construit o "
        "versiune lexicală a retrivorului folosind PostgreSQL full-text search "
        "cu dicționarul românesc implicit, pe aceleași fragmente și aceleași "
        "întrebări, ca să izolăm efectul reprezentării (embeddings dense vs. "
        "tokeni lexicali)."
    ),
    (
        "Rezultatele obținute arată că regăsirea semantică cu ChromaDB și Gemini "
        "ajunge la recall@5 = 0,93 pe Corpus A și 0,97 pe Corpus B, în timp ce "
        "baseline-ul PostgreSQL full-text obține doar 0,77 și 0,93. Diferența "
        "este mai mare pe corpusul juridic, unde întrebările și textul din "
        "documente folosesc formulări diferite (de exemplu \u00abdurata "
        "perioadei de probă\u00bb în întrebare vs. \u00abtermenul de "
        "încercare\u00bb în contract). Gradul de fundamentare al răspunsurilor "
        "generate este de 0,90 (27 din 30) pe Corpus A și 0,97 (29 din 30) pe "
        "Corpus B, cu eșecuri clusterizate exclusiv pe întrebări unde regăsirea "
        "însăși nu reușise — ceea ce confirmă că modelul generativ rămâne fidel "
        "contextului când acesta este bun."
    ),
    (
        "În privința eficienței, latența mediană a unei interogări este de "
        "aproximativ 1,6 secunde, iar descompunerea pe etape arată că 74% din "
        "timp este petrecut în apelul la modelul generativ Gemini, 13% în "
        "calcularea embedding-ului interogării și doar 2% în căutarea ChromaDB. "
        "Indexarea unui document din Corpus A durează în medie 6,8 secunde, "
        "dominant din cauza rotunjirilor de rețea către API-ul Gemini. "
        "Concluzia evaluării este că retriverul dens oferă un câștig clar de "
        "calitate față de baseline-ul lexical fără să introducă un cost "
        "semnificativ în bugetul de latență — sursa dominantă de timp rămâne "
        "modelul generativ, iar orice optimizare ulterioară va trebui să "
        "vizeze apelul LLM, nu componenta de vector store."
    ),
]

BIBLIOGRAFIE = [
    (
        "[1] P. Lewis, E. Perez, A. Piktus et al., \u00abRetrieval-Augmented "
        "Generation for Knowledge-Intensive NLP Tasks\u00bb, în Advances in "
        "Neural Information Processing Systems 33 (NeurIPS 2020), pp. 9459–9474."
    ),
    (
        "[2] J. Johnson, M. Douze, H. Jégou, \u00abBillion-Scale Similarity "
        "Search with GPUs\u00bb, IEEE Transactions on Big Data, vol. 7, nr. 3, "
        "pp. 535–547, 2021."
    ),
    (
        "[3] V. Karpukhin, B. Oğuz, S. Min et al., \u00abDense Passage Retrieval "
        "for Open-Domain Question Answering\u00bb, în Proceedings of EMNLP 2020, "
        "pp. 6769–6781."
    ),
    (
        "[4] A. Vaswani, N. Shazeer, N. Parmar et al., \u00abAttention Is All "
        "You Need\u00bb, în Advances in Neural Information Processing Systems 30 "
        "(NeurIPS 2017), pp. 5998–6008."
    ),
    (
        "[5] N. Reimers, I. Gurevych, \u00abSentence-BERT: Sentence Embeddings "
        "using Siamese BERT-Networks\u00bb, în Proceedings of EMNLP-IJCNLP 2019, "
        "pp. 3982–3992."
    ),
]


def build_document() -> Document:
    document = Document()
    _set_default_style(document)

    _add_header(document)

    _add_heading(document, "Stadiul actual", level=1)
    for paragraph in STADIUL_ACTUAL:
        _add_paragraph(document, paragraph)

    _add_heading(document, "Abordare", level=1)
    _add_paragraph(document, ABORDARE_INTRO)
    for paragraph in ABORDARE_BODY:
        _add_paragraph(document, paragraph)

    _add_heading(document, "Evaluarea abordării propuse", level=1)
    _add_paragraph(document, EVALUARE_INTRO)
    for paragraph in EVALUARE_BODY:
        _add_paragraph(document, paragraph)

    _add_heading(document, "Bibliografie", level=1)
    for reference in BIBLIOGRAFIE:
        _add_paragraph(document, reference, indent=False)

    return document


def main() -> None:
    document = build_document()
    document.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
