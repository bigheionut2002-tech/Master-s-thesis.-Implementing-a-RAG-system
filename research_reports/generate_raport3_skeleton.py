"""Generate the empty skeleton for Raport 3 (.docx).

Run once with:

    cd "/Users/ionut/masters thesis"
    rag_application/backend/.venv/bin/python \
        research_reports/generate_raport3_skeleton.py

The resulting file is ``research_reports/Raport3_RAG_VectorDB.docx``. It
contains only section headings and placeholders — the actual prose will be
written in a later step.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

OUTPUT_PATH = Path(__file__).parent / "Raport3_RAG_VectorDB.docx"

TITLE_RO = (
    "Sisteme de Regăsire a Informațiilor bazate pe Baze de Date Vectoriale. "
    "Implementarea unui sistem RAG"
)
STUDENT = "Bighe Ionuț Denis"
COORDINATOR = "Prof. Dr. Maria-Camelia Chișălița-Crețu"


def _set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def build_document() -> Document:
    document = Document()
    _set_default_font(document)

    document.add_heading("Raport 3 — Proiect de cercetare în baze de date", level=0)

    document.add_paragraph().add_run(f"Titlul proiectului de cercetare: {TITLE_RO}").bold = True
    document.add_paragraph(f"Student: {STUDENT}")
    document.add_paragraph(f"Coordonator științific: {COORDINATOR}")
    document.add_paragraph()

    document.add_heading("Stadiul actual", level=1)
    document.add_paragraph(
        "[Aproximativ ½ pagină. Descrieți contribuțiile principale și cele mai "
        "recente în domeniu și legătura lor cu problema abordată. Exemple: "
        "„Lewis et al. [1] prezintă o soluție pentru …, iar noi propunem o "
        "optimizare a …” sau „Soluția propusă în [2] nu consideră cerințe "
        "speciale pentru …”.]"
    )

    document.add_heading("Abordare", level=1)
    document.add_paragraph(
        "[Aproximativ 1–1½ pagini. Descrieți problema de rezolvat și abordarea "
        "folosită; prezentați contribuțiile principale ale proiectului.]"
    )

    document.add_heading("Evaluarea abordării propuse", level=1)
    document.add_paragraph(
        "[Aproximativ ½–1 pagină. Evaluare experimentală: compararea "
        "rezultatelor obținute cu alte soluții pentru diferite seturi de date; "
        "compararea abordării în termeni de eficiență și performanță față de "
        "alte soluții; utilizarea unor metrici; concluzii ale studiului.]"
    )

    document.add_heading("Bibliografie", level=1)
    document.add_paragraph(
        "[Lista referințelor bibliografice folosite, în format numeric "
        "(ex.: [1] Lewis et al., Retrieval-Augmented Generation for "
        "Knowledge-Intensive NLP Tasks, NeurIPS 2020).]"
    )

    return document


def main() -> None:
    document = build_document()
    document.save(str(OUTPUT_PATH))
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
